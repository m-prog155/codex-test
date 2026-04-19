from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

import pythoncom
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont
import win32com.client as win32
from win32com.client import constants


ROOT = Path(r"D:\Projects\Car")
DOCS = ROOT / "docs"

SENIOR_TEMPLATE = DOCS / "senior-template.docx"
SOURCE_DOC = DOCS / "current-thesis.docx"
OUTPUT_DOC = DOCS / "thesis-senior-filled.docx"
FINAL_DOC = DOCS / "基于深度学习的车辆类型与车牌检测识别系统-论文初稿.docx"
BACKUP_DOC = DOCS / "基于深度学习的车辆类型与车牌检测识别系统-论文初稿.backup-20260418-before-senior-fill.docx"
COVER_BLOCK_IMAGE = ROOT / "outputs" / "thesis_figures" / "cover_info_block.png"

THESIS_TITLE_CN = "基于深度学习的车辆类型与车牌检测识别系统"


def normalize_text(text: str) -> str:
    return (
        text.replace("\r", "")
        .replace("\x07", "")
        .replace("\n", "")
        .replace("\x0b", "")
        .strip()
    )


def nonempty_paragraphs(doc, limit: int | None = None):
    result = []
    for idx in range(1, doc.Paragraphs.Count + 1):
        paragraph = doc.Paragraphs(idx)
        text = normalize_text(paragraph.Range.Text)
        if text:
            result.append(paragraph)
            if limit is not None and len(result) >= limit:
                break
    return result


def find_paragraph(doc, *, exact: str | None = None, startswith: str | None = None):
    for idx in range(1, doc.Paragraphs.Count + 1):
        paragraph = doc.Paragraphs(idx)
        text = normalize_text(paragraph.Range.Text)
        if exact is not None and text == exact:
            return paragraph
        if startswith is not None and text.startswith(startswith):
            return paragraph
    raise RuntimeError(f"paragraph not found: exact={exact!r}, startswith={startswith!r}")


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.Range.Text = text + "\r"


def copy_formatted_range(src_start, src_end, dst_start, dst_end) -> None:
    src_range = src_start.Range.Duplicate
    src_range.End = src_end.Range.End
    dst_range = dst_start.Range.Duplicate
    dst_range.End = dst_end.Range.End
    dst_range.FormattedText = src_range.FormattedText


def fill_cover(dest_doc) -> None:
    cover_lines = [
        "阳 光 学 院",
        "本科毕业论文(设计)",
        "题    目：   基于深度学习的车辆类型与",
        "   车牌检测识别系统",
        "院(系)别：         人工智能学院",
        "专    业：     数据科学与大数据技术专升本",
        "年    级：            2024级",
        "学    号：         2024349243",
        "姓    名：             陈东东",
        "指导教师：     陈志聪教授    陈辉煌老师",
        "2026年     05月",
    ]
    cover_paragraphs = nonempty_paragraphs(dest_doc, limit=11)
    if len(cover_paragraphs) < 11:
        raise RuntimeError("failed to locate senior cover paragraphs")
    for paragraph, text in zip(cover_paragraphs, cover_lines):
        replace_paragraph_text(paragraph, text)


def fill_abstract(dest_doc, src_doc) -> None:
    src_title = find_paragraph(src_doc, exact=THESIS_TITLE_CN)
    src_keywords = find_paragraph(src_doc, startswith="关键词：")
    dst_title = find_paragraph(dest_doc, startswith="基于YOLOv5")
    dst_keywords = find_paragraph(dest_doc, startswith="关键词：")
    copy_formatted_range(src_title, src_keywords, dst_title, dst_keywords)


def fill_body(dest_doc, src_doc) -> None:
    src_body_start = find_paragraph(src_doc, exact="1 绪论")
    dst_body_start = find_paragraph(dest_doc, exact="1 绪论")
    src_range = src_doc.Range(src_body_start.Range.Start, src_doc.Content.End - 1)
    dst_range = dest_doc.Range(dst_body_start.Range.Start, dest_doc.Content.End - 1)
    dst_range.FormattedText = src_range.FormattedText


def update_fields(dest_doc) -> None:
    dest_doc.Fields.Update()
    if dest_doc.TablesOfContents.Count >= 1:
        dest_doc.TablesOfContents(1).Update()
        dest_doc.TablesOfContents(1).UpdatePageNumbers()


def render_cover_block_image() -> None:
    COVER_BLOCK_IMAGE.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1600, 760
    image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)

    label_font = ImageFont.truetype(r"C:\Windows\Fonts\simsun.ttc", 54)
    value_font = ImageFont.truetype(r"C:\Windows\Fonts\STKAITI.TTF", 54)
    suffix_font = ImageFont.truetype(r"C:\Windows\Fonts\simsun.ttc", 48)

    x_label_right = 300
    x_line_start = 360
    x_line_end = 1260
    x_value_center = (x_line_start + x_line_end) / 2
    x_suffix = 1290
    top = 20
    row_h = 110

    rows = [
        ("院(系)别：", "人工智能学院", "", False),
        ("专    业：", "数据科学与大数据技术", "", False),
        ("年    级：", "2024级", "", False),
        ("学    号：", "2024349243", "", False),
        ("姓    名：", "陈东东", "", False),
        ("指导教师：", "", "（签名（职称））", False),
    ]

    for idx, (label, value, suffix, is_title) in enumerate(rows):
        y = top + idx * row_h
        text_y = y + 18
        line_y = y + 78

        if label:
            bbox = draw.textbbox((0, 0), label, font=label_font)
            draw.text((x_label_right - (bbox[2] - bbox[0]), text_y), label, fill="black", font=label_font)

        line_end = x_line_end if not suffix else 1080
        draw.line((x_line_start, line_y, line_end, line_y), fill="black", width=2)

        if value:
            bbox = draw.textbbox((0, 0), value, font=value_font)
            tx = x_value_center - (bbox[2] - bbox[0]) / 2
            draw.text((tx, text_y), value, fill="black", font=value_font)

        if suffix:
            draw.text((x_suffix, text_y), suffix, fill="black", font=suffix_font)

    image.save(COVER_BLOCK_IMAGE)


def apply_cover_info_table(doc) -> None:
    title_second_para = find_paragraph(doc, startswith="与车牌检测识别系统")
    start_para = find_paragraph(doc, startswith="院(系)别")
    end_para = find_paragraph(doc, startswith="指导教师")
    start = start_para.Range.Start
    end = end_para.Range.End

    rng = doc.Range(start, end)
    rng.Delete()

    insert_range = title_second_para.Range.Duplicate
    insert_range.Collapse(constants.wdCollapseEnd)
    insert_range.InsertParagraphAfter()
    insert_range.Collapse(constants.wdCollapseEnd)

    table = doc.Tables.Add(Range=insert_range, NumRows=6, NumColumns=3)
    table.AllowAutoFit = False
    table.Rows.Alignment = constants.wdAlignRowCenter
    table.Borders.Enable = False
    table.TopPadding = 0
    table.BottomPadding = 0
    table.LeftPadding = 0
    table.RightPadding = 0

    app = doc.Application
    table.Columns(1).Width = app.CentimetersToPoints(3.2)
    table.Columns(2).Width = app.CentimetersToPoints(8.8)
    table.Columns(3).Width = app.CentimetersToPoints(3.8)

    rows = [
        ("院(系)别：", "人工智能学院", ""),
        ("专    业：", "数据科学与大数据技术", ""),
        ("年    级：", "2024级", ""),
        ("学    号：", "2024349243", ""),
        ("姓    名：", "陈东东", ""),
        ("指导教师：", "", "（签名（职称））"),
    ]

    for idx, (label, value, suffix) in enumerate(rows, start=1):
        row = table.Rows(idx)
        row.HeightRule = constants.wdRowHeightAtLeast
        row.Height = app.CentimetersToPoints(0.92)

        label_cell = table.Cell(idx, 1)
        value_cell = table.Cell(idx, 2)
        suffix_cell = table.Cell(idx, 3)

        for cell in (label_cell, value_cell, suffix_cell):
            cell.VerticalAlignment = constants.wdCellAlignVerticalCenter
            cell.TopPadding = 0
            cell.BottomPadding = 0
            cell.LeftPadding = 0
            cell.RightPadding = 0
            cell.Range.ParagraphFormat.SpaceBefore = 0
            cell.Range.ParagraphFormat.SpaceAfter = 0
            cell.Range.ParagraphFormat.LineSpacingRule = constants.wdLineSpaceSingle
            cell.Range.ParagraphFormat.LeftIndent = 0
            cell.Range.ParagraphFormat.RightIndent = 0
            for border_idx in [
                constants.wdBorderLeft,
                constants.wdBorderTop,
                constants.wdBorderRight,
                constants.wdBorderBottom,
            ]:
                cell.Borders(border_idx).LineStyle = constants.wdLineStyleNone

        label_cell.Range.Text = label
        label_cell.Range.Font.Name = "宋体"
        label_cell.Range.Font.NameFarEast = "宋体"
        label_cell.Range.Font.Size = 18
        label_cell.Range.ParagraphFormat.Alignment = constants.wdAlignParagraphRight

        value_cell.Range.Text = value
        value_cell.Range.Font.Name = "楷体"
        value_cell.Range.Font.NameFarEast = "楷体"
        value_cell.Range.Font.Size = 18
        value_cell.Range.ParagraphFormat.Alignment = constants.wdAlignParagraphCenter
        value_cell.Borders(constants.wdBorderBottom).LineStyle = constants.wdLineStyleSingle
        value_cell.Borders(constants.wdBorderBottom).LineWidth = constants.wdLineWidth050pt

        suffix_cell.Range.Text = suffix
        suffix_cell.Range.Font.Name = "宋体"
        suffix_cell.Range.Font.NameFarEast = "宋体"
        suffix_cell.Range.Font.Size = 18
        suffix_cell.Range.ParagraphFormat.Alignment = constants.wdAlignParagraphLeft


def replace_cover_table_with_image(doc) -> None:
    render_cover_block_image()
    if doc.Tables.Count < 1:
        return

    table = doc.Tables(1)
    insert_at = table.Range.Start
    table.Delete()
    rng = doc.Range(insert_at, insert_at)
    rng.ParagraphFormat.Alignment = constants.wdAlignParagraphCenter
    shape = doc.InlineShapes.AddPicture(
        FileName=str(COVER_BLOCK_IMAGE),
        LinkToFile=False,
        SaveWithDocument=True,
        Range=rng,
    )
    shape.LockAspectRatio = True
    shape.Width = doc.Application.CentimetersToPoints(14.5)


def finalize_cover_block_with_python(doc_path: Path) -> None:
    render_cover_block_image()
    doc = Document(doc_path)
    if doc.tables:
        first_table = doc.tables[0]
        first_table._element.getparent().remove(first_table._element)

    start_idx = None
    end_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if start_idx is None and text.startswith("院(系)别"):
            start_idx = idx
        if text.startswith("指导教师"):
            end_idx = idx
            break

    if start_idx is not None and end_idx is not None and end_idx >= start_idx:
        for idx in range(end_idx, start_idx - 1, -1):
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)

    date_para = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("2026") or text.startswith("2026 年") or "2026" in text and "05" in text:
            date_para = paragraph
            break
    if date_para is None:
        raise RuntimeError("date paragraph not found when finalizing cover")

    image_para = date_para.insert_paragraph_before()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_before = Pt(24)
    image_para.paragraph_format.space_after = Pt(12)
    run = image_para.add_run()
    run.add_picture(str(COVER_BLOCK_IMAGE), width=Cm(14.5))
    doc.save(doc_path)


def postprocess_cover(doc_path: Path) -> None:
    doc = Document(doc_path)
    senior = Document(SENIOR_TEMPLATE)

    def find_para(document: Document, *, exact: str | None = None, startswith: str | None = None):
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if exact is not None and text == exact:
                return paragraph
            if startswith is not None and text.startswith(startswith):
                return paragraph
        raise RuntimeError(f"cover paragraph not found: exact={exact!r}, startswith={startswith!r}")

    def clone_para(dst_para, src_para):
        dst_para._element.getparent().replace(dst_para._element, deepcopy(src_para._element))

    school_dst = find_para(doc, exact="阳 光 学 院")
    school_src = find_para(senior, exact="阳 光 学 院")
    paper_dst = find_para(doc, exact="本科毕业论文(设计)")
    paper_src = find_para(senior, exact="本科毕业论文(设计)")
    title1_dst = find_para(doc, startswith="题")
    title1_src = find_para(senior, startswith="题")
    title2_dst = find_para(doc, startswith="车牌检测识别系统")
    title2_src = find_para(senior, startswith="员检测算法的研究与应用")
    date_dst = find_para(doc, startswith="2026年")
    date_src = find_para(senior, startswith="2024年")

    clone_para(school_dst, school_src)
    clone_para(paper_dst, paper_src)
    clone_para(title1_dst, title1_src)
    clone_para(title2_dst, title2_src)
    clone_para(date_dst, date_src)

    def set_run_texts(paragraph, texts: list[str]) -> None:
        p = paragraph
        for run, text in zip(p.runs, texts):
            run.text = text
        for run in p.runs[len(texts) :]:
            run.text = ""

    set_run_texts(find_para(doc, exact="阳 光 学 院"), ["阳 光 学 院"])
    set_run_texts(find_para(doc, exact="本科毕业论文(设计)"), ["本科毕业", "论文", "(", "设计", ")"])
    set_run_texts(find_para(doc, startswith="题"), ["题", "    ", "目：", " ", "  ", "基于深度学习的车辆类型", "  ", " "])
    set_run_texts(find_para(doc, startswith="员检测算法的研究与应用"), ["  ", " ", "与车牌检测识别系统", "", "", "  ", "  "])
    set_run_texts(find_para(doc, startswith="2024年"), ["               2026", "年", "     05", "月", "", ""])

    doc.save(doc_path)


def main() -> None:
    if not SENIOR_TEMPLATE.exists():
        raise FileNotFoundError(SENIOR_TEMPLATE)
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(SOURCE_DOC)

    if FINAL_DOC.exists():
        shutil.copy2(FINAL_DOC, BACKUP_DOC)

    shutil.copy2(SENIOR_TEMPLATE, OUTPUT_DOC)

    pythoncom.CoInitialize()
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    src = None
    dest = None
    try:
        src = word.Documents.Open(str(SOURCE_DOC))
        dest = word.Documents.Open(str(OUTPUT_DOC))

        fill_cover(dest)
        fill_abstract(dest, src)
        fill_body(dest, src)

        dest.Save()
        dest.Close(False)
        src.Close(False)
        word.Quit()
    finally:
        try:
            if dest is not None:
                dest.Close(False)
        except Exception:
            pass
        try:
            if src is not None:
                src.Close(False)
        except Exception:
            pass
        try:
            word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()

    postprocess_cover(OUTPUT_DOC)

    pythoncom.CoInitialize()
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = None
    try:
        doc = word.Documents.Open(str(OUTPUT_DOC))
        apply_cover_info_table(doc)
        update_fields(doc)
        doc.Save()
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()

    finalize_cover_block_with_python(OUTPUT_DOC)

    shutil.copy2(OUTPUT_DOC, FINAL_DOC)


if __name__ == "__main__":
    main()
