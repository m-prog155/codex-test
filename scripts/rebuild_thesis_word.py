from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(r"D:\Projects\Car")
DOCS = ROOT / "docs"
EXPORT_MD = DOCS / "thesis-word-export-source.md"
REFERENCE_DOC = DOCS / "senior-template.docx"
OUTPUT_DOC = DOCS / "基于深度学习的车辆类型与车牌检测识别系统-论文初稿.docx"
TEMP_DOC = DOCS / "thesis-draft-test.docx"
BACKUP_DOC = DOCS / "基于深度学习的车辆类型与车牌检测识别系统-论文初稿.backup-20260418-before-rebuild.docx"
PANDOC = Path(r"C:\Program Files\Pandoc\pandoc.exe")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def para_text(elem: ET.Element) -> str:
    return "".join(t.text or "" for t in elem.findall(".//w:t", NS)).strip()


def make_page_break() -> ET.Element:
    p = ET.Element(w("p"))
    r = ET.SubElement(p, w("r"))
    br = ET.SubElement(r, w("br"))
    br.set(w("type"), "page")
    return p


def make_toc_para(text: str, style: str) -> ET.Element:
    p = ET.Element(w("p"))
    p_pr = ET.SubElement(p, w("pPr"))
    p_style = ET.SubElement(p_pr, w("pStyle"))
    p_style.set(w("val"), style)
    r = ET.SubElement(p, w("r"))
    t = ET.SubElement(r, w("t"))
    t.text = text
    return p


def build_static_toc() -> list[ET.Element]:
    entries = [
        ("ab", "目    录"),
        ("TOC1", "1 绪论\t1"),
        ("TOC2", "1.1 课题研究背景与意义\t1"),
        ("TOC2", "1.2 国内外研究现状概述\t1"),
        ("TOC2", "1.3 本文研究内容\t2"),
        ("TOC2", "1.4 论文结构\t2"),
        ("TOC1", "2 国内外研究现状与相关技术分析\t3"),
        ("TOC2", "2.1 车辆目标检测研究现状\t3"),
        ("TOC2", "2.2 车牌检测与识别研究现状\t4"),
        ("TOC2", "2.3 本文相关技术路线分析\t4"),
        ("TOC1", "3 系统设计与实现\t6"),
        ("TOC2", "3.1 系统总体设计\t6"),
        ("TOC2", "3.2 系统架构设计\t6"),
        ("TOC2", "3.3 核心功能模块实现\t7"),
        ("TOC2", "3.4 系统运行流程\t8"),
        ("TOC1", "4 实验结果与系统效果分析\t10"),
        ("TOC2", "4.1 实验环境与参数设置\t10"),
        ("TOC2", "4.2 车牌检测模型训练结果分析\t10"),
        ("TOC2", "4.3 系统样例运行效果分析\t11"),
        ("TOC2", "4.4 OCR 小样本评估结果分析\t12"),
        ("TOC2", "4.5 车牌专用 PaddleOCR 对比评估\t13"),
        ("TOC1", "5 总结与展望\t15"),
        ("TOC2", "5.1 全文工作总结\t15"),
        ("TOC2", "5.2 研究工作的不足\t15"),
        ("TOC2", "5.3 后续展望\t16"),
        ("TOC1", "参考文献\t17"),
    ]
    return [make_toc_para(text, style) for style, text in entries]


def find_style(doc: Document, style_key: str):
    for style in doc.styles:
        if style.style_id == style_key or style.name == style_key:
            return style
    return None


def set_run_font(run, font_name: str | None = None, size_pt: float | None = None, bold: bool | None = None) -> None:
    if font_name:
        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_para_text(paragraph, text: str) -> None:
    paragraph.text = text
    if not paragraph.runs:
        paragraph.add_run("")


def set_para_style(doc: Document, paragraph, style_key: str) -> None:
    style = find_style(doc, style_key)
    if style is not None:
        paragraph.style = style


def has_drawing(paragraph) -> bool:
    return any(elem.tag == qn("w:drawing") for elem in paragraph._element.iter())


def format_cover(doc: Document) -> None:
    cover = doc.paragraphs[:11]
    if len(cover) < 11:
        return

    cover_lines = [
        "阳 光 学 院",
        "本科毕业论文(设计)",
        "题    目：   基于深度学习的车辆类型与",
        "   车牌检测识别系统",
        "院(系)别：         人工智能学院",
        "专    业：     数据科学与大数据技术专升本",
        "年    级：            2024级",
        "学    号：         2024349243",
        "姓    名：             陈东东",
        "指导教师：     陈志聪教授    陈辉煌老师",
        "               2026年5月",
    ]

    for paragraph, text in zip(cover, cover_lines):
        set_para_style(doc, paragraph, "a")
        set_para_text(paragraph, text)
        paragraph.alignment = None
        pf = paragraph.paragraph_format
        pf.space_before = None
        pf.space_after = None
        pf.line_spacing = None
        pf.first_line_indent = None
        pf.left_indent = None

    cover[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover[0].paragraph_format.space_before = Pt(16.6)
    cover[0].paragraph_format.space_after = Pt(8.3)
    set_run_font(cover[0].runs[0], "华文行楷", 42)

    cover[1].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    cover[1].paragraph_format.space_before = Pt(16.6)
    cover[1].paragraph_format.space_after = Pt(8.3)
    cover[1].paragraph_format.left_indent = Pt(36)
    set_run_font(cover[1].runs[0], "黑体", 31.5)

    cover[2].paragraph_format.first_line_indent = Pt(54)
    cover[2].paragraph_format.line_spacing = Pt(38)
    set_run_font(cover[2].runs[0], "楷体", 18)

    cover[3].paragraph_format.first_line_indent = Pt(144)
    cover[3].paragraph_format.line_spacing = Pt(38)
    set_run_font(cover[3].runs[0], "楷体", 18)

    for paragraph in cover[4:10]:
        paragraph.paragraph_format.first_line_indent = Pt(54)
        paragraph.paragraph_format.line_spacing = Pt(39)
        set_run_font(paragraph.runs[0], "楷体", 18)

    cover[10].paragraph_format.first_line_indent = Pt(36)
    set_run_font(cover[10].runs[0], "楷体", 18)


def format_abstracts(doc: Document) -> None:
    paragraphs = doc.paragraphs
    cn_title = next((p for p in paragraphs if p.text.strip() == "基于深度学习的车辆类型与车牌检测识别系统"), None)
    cn_heading = next((p for p in paragraphs if p.text.strip() in {"摘 要", "摘    要", "摘要"}), None)
    en_title = next(
        (
            p
            for p in paragraphs
            if p.text.strip()
            == "Design and Implementation of a Deep-Learning-Based Vehicle Type and License Plate Detection and Recognition System"
        ),
        None,
    )
    en_heading = next((p for p in paragraphs if p.text.strip() == "Abstract"), None)
    cn_keywords = next((p for p in paragraphs if p.text.strip().startswith("关键词：")), None)
    en_keywords = next((p for p in paragraphs if p.text.strip().startswith("Keywords:")), None)

    if cn_title is not None:
        set_para_style(doc, cn_title, "aa")
    if en_title is not None:
        set_para_style(doc, en_title, "aa")

    for heading in (cn_heading, en_heading):
        if heading is None:
            continue
        set_para_style(doc, heading, "a")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(15.6)
        heading.paragraph_format.space_after = Pt(15.6)
        heading.paragraph_format.line_spacing = Pt(24)
        for run in heading.runs:
            set_run_font(run, "黑体", 14)

    if cn_heading is not None and cn_keywords is not None:
        started = False
        for paragraph in paragraphs:
            if paragraph is cn_heading:
                started = True
                continue
            if not started:
                continue
            if paragraph is cn_keywords:
                break
            if not paragraph.text.strip():
                continue
            paragraph.paragraph_format.first_line_indent = Pt(24)
            if paragraph == paragraphs[paragraphs.index(cn_heading) + 1]:
                set_para_style(doc, paragraph, "a")
            else:
                set_para_style(doc, paragraph, "a9")

    if en_heading is not None and en_keywords is not None:
        between = False
        first_en_body = True
        for paragraph in paragraphs:
            if paragraph is en_heading:
                between = True
                continue
            if not between:
                continue
            if paragraph is en_keywords:
                break
            if not paragraph.text.strip():
                continue
            paragraph.paragraph_format.first_line_indent = Pt(24)
            set_para_style(doc, paragraph, "a" if first_en_body else "a9")
            first_en_body = False


def format_toc_and_captions(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "目    录":
            set_para_style(doc, paragraph, "ab")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith(("图 ", "表 ")):
            set_para_style(doc, paragraph, "af0")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def remove_duplicate_image_labels(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    to_remove = []
    for idx in range(1, len(paragraphs) - 1):
        previous = paragraphs[idx - 1]
        current = paragraphs[idx]
        following = paragraphs[idx + 1]
        current_text = current.text.strip()
        following_text = following.text.strip()
        if not current_text:
            continue
        if has_drawing(previous) and following_text.startswith("图 ") and current_text in following_text:
            to_remove.append(current._element)

    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)


def format_tables(doc: Document) -> None:
    style = find_style(doc, "a7")
    for table in doc.tables:
        if style is not None:
            table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


def resize_images(doc: Document) -> None:
    max_width = Inches(5.6)
    max_height = Inches(7.2)
    for shape in doc.inline_shapes:
        width = shape.width
        height = shape.height
        if width == 0 or height == 0:
            continue
        scale = min(max_width / width, max_height / height, 1.0)
        if scale < 1.0:
            shape.width = int(width * scale)
            shape.height = int(height * scale)


def postprocess_docx() -> None:
    doc = Document(OUTPUT_DOC)
    format_cover(doc)
    format_abstracts(doc)
    remove_duplicate_image_labels(doc)
    format_toc_and_captions(doc)
    format_tables(doc)
    resize_images(doc)
    doc.save(OUTPUT_DOC)


def rebuild_docx() -> None:
    if OUTPUT_DOC.exists():
        shutil.copy2(OUTPUT_DOC, BACKUP_DOC)

    subprocess.run(
        [
            str(PANDOC),
            str(EXPORT_MD.name),
            "-o",
            str(TEMP_DOC.name),
            f"--reference-doc={REFERENCE_DOC}",
        ],
        cwd=DOCS,
        check=True,
    )

    with zipfile.ZipFile(TEMP_DOC, "r") as zin:
        xml_bytes = zin.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        body = root.find("w:body", NS)
        if body is None:
            raise RuntimeError("document.xml missing body")

        date_para = None
        cn_kw_para = None
        en_kw_para = None
        for child in list(body):
            if child.tag != w("p"):
                continue
            text = para_text(child)
            if text == "2026年5月":
                date_para = child
            elif text.startswith("关键词："):
                cn_kw_para = child
            elif text.startswith("Keywords:"):
                en_kw_para = child

        if date_para is None or cn_kw_para is None or en_kw_para is None:
            raise RuntimeError("failed to locate expected thesis paragraphs")

        def insert_after(anchor: ET.Element, new_nodes: list[ET.Element]) -> None:
            idx = list(body).index(anchor)
            for offset, node in enumerate(new_nodes, start=1):
                body.insert(idx + offset, node)

        insert_after(date_para, [make_page_break()])
        insert_after(cn_kw_para, [make_page_break()])
        insert_after(en_kw_para, [make_page_break(), *build_static_toc(), make_page_break()])

        ET.register_namespace("w", W_NS)
        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        fd, temp_zip_path = tempfile.mkstemp(suffix=".docx", dir=str(DOCS))
        os.close(fd)
        Path(temp_zip_path).unlink(missing_ok=True)
        temp_zip = Path(temp_zip_path)
        with zipfile.ZipFile(TEMP_DOC, "r") as zin, zipfile.ZipFile(temp_zip, "w") as zout:
            for item in zin.infolist():
                data = new_xml if item.filename == "word/document.xml" else zin.read(item.filename)
                zout.writestr(item, data)

        shutil.move(temp_zip, OUTPUT_DOC)
    postprocess_docx()


if __name__ == "__main__":
    rebuild_docx()
